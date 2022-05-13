#!/usr/bin/env python
""" plot.py
"""
import sys
import os
import matplotlib as mpl
import matplotlib.pyplot as plt
import pandas as pd
import glob
import re
import docx
from docx import Document
from docx.shared import Inches,Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from jtop import jp

def numerical_sort(value):
    numbers = re.compile(r'(\d+)')
    parts = numbers.split(value)
    parts[1::2] = map(int, parts[1::2])
    return parts

VKEY=str(sys.argv[1])  #sk_o_h
VKEY3=str(sys.argv[2])  #skh
VF2=str(sys.argv[3])  #z\m1.isw.csv
VN2,VEXT=os.path.splitext(os.path.basename(VF2))
VFDOC=str(sys.argv[4])  #z_plot_compare_
VUNIT=int(sys.argv[5])  #1:kN,m 2:KN*m,rad
VFPLAXIS='z_m1_plaxis.csv'
VRSKIP=6
VRSKIP_PLAXIS=1
VCDEPTH=13
VDPI=200
mpl.font_manager._rebuild()
mpl.rc('text', usetex=True)
plt.rcParams.update({'text.latex.preamble' : [], 'font.family' : 'sans-serif'})
df2=pd.read_csv(VF2, header=None, sep=',', skiprows=VRSKIP, engine='python')
df0=pd.read_csv(VFPLAXIS, header=None, sep=',', skiprows=VRSKIP_PLAXIS, engine='python')

VFS=sorted(glob.glob(f'{VKEY}*.txt'), key=numerical_sort)
VFS3=sorted(glob.glob(f'{VKEY3}*.txt'), key=numerical_sort)
for VI in range(len(VFS)):
    VF=VFS[VI]
    VN,VEXT=os.path.splitext(os.path.basename(VF))
    df=pd.read_csv(VF, header=None, sep=r'\t', names=['x','y'], engine='python')
    x=df['x']
    y=df['y']
    VCX=2*VI+1
    VCY=VCX+1
    x2=df2.iloc[:,VCX]
    y2=df2.iloc[:,VCY]*0.001
    if abs(x.min(axis=0))>abs(x.max(axis=0)):
        x=-x
        x2=-x2
    if abs(y.min(axis=0))>abs(y.max(axis=0)):
        y=-y
        y2=-y2
    if VUNIT==1:
        x2=-x2
    elif VUNIT==2:
        y2=-y2
    else:
        raise SystemExit(f'ERROR: invalid VUNIT = {VUNIT}')
    VF3=VFS3[VI]
    VN3,VEXT3=os.path.splitext(os.path.basename(VF3))
    df3=pd.read_csv(VF3, header=None, sep=r'\t', names=['x','y'], engine='python')
    x3=df3['x']
    y3=df3['y']
    VOUT='z_'+str(VN3)+'.png'
    if VF==VFS[-1]:
        VLTITLE='杭下端'
    else:
        VLTITLE='GL'+str('{0:.2f}'.format(df0.iat[VI,VCDEPTH]))+'m'
    plt.tick_params(direction='in')
    plt.plot(x, y, color='blue', label=f'PLAXIS')
    plt.plot(x2, y2, color='yellow', label='pidac')
    plt.plot(x3, y3, color='red', label=f'{str(VN3)}')
    if VUNIT==1:
        plt.xlabel(r'変位(m)')
        plt.ylabel(r'荷重(kN)')
    elif VUNIT==2:
        plt.xlabel(r'回転変位(rad)')
        plt.ylabel(r'モーメント(kN・m)')
    else:
        raise SystemExit(f'ERROR: invalid VUNIT = {VUNIT}')
    plt.xlim(None, None)
    plt.ylim(None, None)
    plt.grid(axis='both')
    plt.legend(framealpha=1.0, title=VLTITLE)  #plt.legend()
    plt.savefig(VOUT, dpi=VDPI)
    plt.clf()
    #ii_svg2emf ${VOUT}

document = Document()
sections = document.sections
for section in sections:
    section.top_margin = Mm(15.0)
    section.bottom_margin = Mm(15.0)
    section.left_margin = Mm(30.0)
    section.right_margin = Mm(30.0)
p = document.add_heading('地盤ばね平滑化前後比較', level=0)
p.runs[0].font.size = docx.shared.Pt(10)
p.runs[0].bold = True
VFS=sorted(glob.glob(f'z_{VKEY3}*.png'), key=numerical_sort)
VI=0
for VF in VFS:
    if VI==0:
        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(str(VF), width=Mm(75.0))
        VI=1
    else:
        run_2 = paragraph.add_run()
        run_2.add_picture(str(VF), width=Mm(75.0))
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        VI=0
VOUT=str(VFDOC)+str(VKEY3)+'.docx'
document.save(VOUT)
