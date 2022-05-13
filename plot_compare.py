#!/usr/bin/env python
""" plot.py
"""
import sys
import os
import matplotlib as mpl
import matplotlib.pyplot as plt
import pandas as pd
import glob
import re
import docx
from docx import Document
from docx.shared import Inches,Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from jtop import jp

def numerical_sort(value):
    numbers = re.compile(r'(\d+)')
    parts = numbers.split(value)
    parts[1::2] = map(int, parts[1::2])
    return parts

VKEY=str(sys.argv[1])  #sk_o_h
VF2=str(sys.argv[2])  #z\m1.isw.csv
VN2,VEXT=os.path.splitext(os.path.basename(VF2))
VFDOC=str(sys.argv[3])  #z_plot_compare_
VFDISP='inputdisp.txt'
VFFORCE='m1result_curve.csv'
VFPLAXIS='z_m1_plaxis.csv'
VRSKIP=6
VRHEAD=2
VRSKIP_DISP=2
VRSKIP_PLAXIS=1
VCDEPTH=13
VDPI=200
#for font in mpl.font_manager.findSystemFonts():
#    print(mpl.font_manager.FontProperties(fname=font).get_name())
#sys.exit()  #test
mpl.font_manager._rebuild()
mpl.rc('text', usetex=True)
plt.rcParams.update({'text.latex.preamble' : [], 'font.family' : 'sans-serif'})
df2=pd.read_csv(VF2, header=None, sep=',', skiprows=VRSKIP, engine='python')
df3=pd.read_csv(VFPLAXIS, header=None, sep=',', skiprows=VRSKIP_PLAXIS, engine='python')

VOUT='z_sk_o_h0_px.png'
VLTITLE='荷重-変位関係'
VCX=1
VCY=2
dfx=pd.read_csv(VFDISP, header=None, sep=r'\t', names=['x'], skiprows=VRSKIP_DISP, engine='python')
dfy=pd.read_csv(VFFORCE, header=None, sep=r'\t', names=['y'], engine='python')
x=dfx['x']
y=dfy['y']
x2=df2.iloc[:,VCX]
y2=df2.iloc[:,VCY]*0.001
plt.tick_params(direction='in')
plt.plot(x, y, color='blue', label='FEM')
plt.plot(x2, y2, color='red', label='Penzien')
plt.xlabel(r'加力点変位 (m)')
plt.ylabel(r'荷重 (kN)')
plt.xlim(None, None)
plt.ylim(None, None)
plt.grid(axis='both')
plt.legend(framealpha=1.0, title=VLTITLE)  #plt.legend()
plt.savefig(VOUT, dpi=VDPI)
plt.clf()

VFS=sorted(glob.glob(f'{VKEY}*.txt'), key=numerical_sort)
#print(VFS)  #test
#print(df2)  #test
#print(VLTITLES)  #test
#sys.exit()  #test
VI=0
for VF in VFS:  #compare sk_o_h
    VN,VEXT=os.path.splitext(os.path.basename(VF))
    VOUT='z_'+str(VN)+'.png'
    if VF==VFS[-1]:
        VLTITLE='杭下端'
    else:
        VLTITLE='GL'+str('{0:.2f}'.format(df3.iloc[VI,VCDEPTH]))+'m'
    VCY=VCY+1
    dfy=pd.read_csv(VF, header=None, sep=r'\t', names=['y','dummy'], engine='python')
    x=dfx['x']
    y=dfy['y']
    x2=df2.iloc[:,[VCX]]
    y2=df2.iloc[:,[VCY]]
    plt.tick_params(direction='in')
    plt.plot(x, y, color='blue', label='FEM')
    plt.plot(x2, y2, color='red', label='Penzien')
    plt.xlabel(r'加力点変位 (m)')
    plt.ylabel(r'変位 (m)')
    plt.xlim(None, None)
    plt.ylim(None, None)
    plt.grid(axis='both')
    plt.legend(framealpha=1.0, title=VLTITLE)  #plt.legend()
    plt.savefig(VOUT, dpi=VDPI)
    plt.clf()
    VI=VI+1
    #ii_svg2emf ${VOUT}

document = Document()
sections = document.sections
for section in sections:
    section.top_margin = Mm(15.0)
    section.bottom_margin = Mm(15.0)
    section.left_margin = Mm(30.0)
    section.right_margin = Mm(30.0)
p = document.add_heading('FEM - Penzien 解析結果比較', level=0)
p.runs[0].font.size = docx.shared.Pt(10)
p.runs[0].bold = True
VFS=sorted(glob.glob(f'z_{VKEY}*.png'), key=numerical_sort)
VI=0
for VF in VFS:
    if VI==0:
        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(str(VF), width=Mm(75.0))
        VI=1
    else:
        run_2 = paragraph.add_run()
        run_2.add_picture(str(VF), width=Mm(75.0))
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        VI=0
VOUT=str(VFDOC)+str(VKEY)+'.docx'
document.save(VOUT)
